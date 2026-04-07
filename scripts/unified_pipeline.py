from __future__ import annotations

import argparse
import logging
import re
import time
from pathlib import Path
from typing import Any, Dict

from docling.datamodel.accelerator_options import AcceleratorDevice, AcceleratorOptions
from docling.datamodel.base_models import InputFormat
from docling.datamodel.pipeline_options import (
    EasyOcrOptions,
    OcrMacOptions,
    PdfPipelineOptions,
    RapidOcrOptions,
    TableStructureOptions,
)
from docling.document_converter import (
    DocumentConverter,
    ExcelFormatOption,
    PdfFormatOption,
    PowerpointFormatOption,
    WordFormatOption,
)
from docling_core.types.doc import ImageRefMode
from docling_core.types.doc.document import DoclingDocument
from docx import Document
from docx.shared import Inches


LOG = logging.getLogger("unified_pipeline")
IMAGE_TAG_PATTERN = re.compile(r"!\[(?P<alt>[^\]]*)\]\((?P<path>[^)]+)\)")


def configure_logging(verbose: bool) -> None:
    level = logging.DEBUG if verbose else logging.INFO
    logging.basicConfig(
        level=level,
        format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    )


def build_ocr_options(ocr_engine: str):
    if ocr_engine == "none":
        return None
    if ocr_engine == "rapidocr":
        return RapidOcrOptions()
    if ocr_engine == "ocrmac":
        return OcrMacOptions()
    if ocr_engine == "easyocr":
        return EasyOcrOptions(download_enabled=False)
    if ocr_engine == "auto":
        return RapidOcrOptions()
    return None


def detect_input_format(input_file: Path) -> InputFormat:
    suffix = input_file.suffix.lower()
    if suffix == ".pdf":
        return InputFormat.PDF
    if suffix == ".docx":
        return InputFormat.DOCX
    if suffix == ".pptx":
        return InputFormat.PPTX
    if suffix == ".xlsx":
        return InputFormat.XLSX
    if suffix == ".doc":
        raise ValueError(
            "Legacy .doc files are not supported directly. Convert the file to .docx first."
        )
    raise ValueError(
        f"Unsupported file type: {suffix}. Supported types are .pdf, .docx, .pptx, and .xlsx."
    )


def build_converter(args: argparse.Namespace) -> DocumentConverter:
    input_format = detect_input_format(args.input_file)

    if input_format == InputFormat.PDF:
        accelerator_options = AcceleratorOptions(device=AcceleratorDevice(args.device))

        pipeline_options = PdfPipelineOptions(
            artifacts_path=str(args.artifacts_path),
            enable_remote_services=False,
            accelerator_options=accelerator_options,
            do_ocr=args.ocr_engine != "none",
            do_table_structure=not args.no_tables,
            generate_picture_images=True,
        )
        pipeline_options.table_structure_options = TableStructureOptions()
        pipeline_options.table_structure_options.do_cell_matching = True
        pipeline_options.images_scale = args.image_scale

        ocr_options = build_ocr_options(args.ocr_engine)
        if ocr_options:
            pipeline_options.ocr_options = ocr_options

        format_option = PdfFormatOption(pipeline_options=pipeline_options)
    elif input_format == InputFormat.DOCX:
        format_option = WordFormatOption()
    elif input_format == InputFormat.PPTX:
        format_option = PowerpointFormatOption()
    elif input_format == InputFormat.XLSX:
        format_option = ExcelFormatOption()
    else:
        raise ValueError(f"Unsupported input format: {input_format}")

    return DocumentConverter(format_options={input_format: format_option})


def run_docling(args: argparse.Namespace) -> tuple[DoclingDocument, Path, Path]:
    converter = build_converter(args)
    result = converter.convert(str(args.input_file))
    doc = result.document

    args.output_dir.mkdir(parents=True, exist_ok=True)
    image_dir = args.output_dir / "images"
    stem = args.input_file.stem
    markdown_path = args.output_dir / f"{stem}_read_llm.md"
    json_path = args.output_dir / f"{stem}_document.json"

    LOG.info("Saving JSON to %s", json_path)
    doc.save_as_json(
        json_path,
        artifacts_dir=image_dir,
        image_mode=ImageRefMode.REFERENCED,
        indent=2,
        coord_precision=2,
    )

    LOG.info("Saving LLM-ready markdown to %s", markdown_path)
    doc.save_as_markdown(
        markdown_path,
        artifacts_dir=image_dir,
        image_mode=ImageRefMode.REFERENCED,
    )
    return doc, markdown_path, image_dir


def build_llm_stub(markdown_path: Path) -> Dict[str, Any]:
    prompt_preview = markdown_path.read_text(encoding="utf-8")[:800]
    return {
        "message": "LLM call not executed (stub). Plug in OpenAI or Groq here.",
        "prompt_preview": prompt_preview,
        "markdown_path": str(markdown_path),
    }


def describe_image_with_llm(image_path: Path) -> str:
    # Replace this stub with the real OpenAI vision call later.
    return f"Stub response for {image_path.name}. Replace describe_image_with_llm() with the OpenAI API call."


def augment_markdown_with_llm(markdown_path: Path) -> tuple[Path, list[dict[str, str]]]:
    markdown = markdown_path.read_text(encoding="utf-8")
    inserted_notes: list[dict[str, str]] = []
    markdown_dir = markdown_path.parent

    def replace(match: re.Match[str]) -> str:
        alt_text = match.group("alt")
        raw_image_path = match.group("path").strip()
        resolved_image_path = (markdown_dir / raw_image_path).resolve()

        try:
            if not resolved_image_path.exists():
                description = f"Image description unavailable: file not found at {resolved_image_path}"
            else:
                description = describe_image_with_llm(resolved_image_path)
        except Exception as exc:  # noqa: BLE001
            LOG.warning("Image description failed for %s: %s", resolved_image_path, exc)
            description = f"Image description unavailable: {exc}"

        inserted_notes.append(
            {
                "alt": alt_text,
                "image_path": str(resolved_image_path),
                "description": description,
            }
        )
        return f"{match.group(0)}\n\n*LLM image note:* {description}\n"

    augmented_markdown = IMAGE_TAG_PATTERN.sub(replace, markdown)
    augmented_path = markdown_path.with_name("read_llm_augmented.md")
    augmented_path.write_text(augmented_markdown, encoding="utf-8")
    LOG.info("Saved augmented markdown to %s", augmented_path)
    return augmented_path, inserted_notes


def generate_docx_from_augmented_markdown(
    markdown_path: Path, output_docx: Path, title: str | None = None
) -> None:
    docx_doc = Document()
    if title:
        docx_doc.add_heading(title, level=1)

    markdown_dir = markdown_path.parent
    lines = markdown_path.read_text(encoding="utf-8").splitlines()

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue

        image_match = IMAGE_TAG_PATTERN.fullmatch(line)
        if image_match:
            resolved_image_path = (markdown_dir / image_match.group("path").strip()).resolve()
            if resolved_image_path.exists():
                docx_doc.add_picture(str(resolved_image_path), width=Inches(5.5))
            else:
                docx_doc.add_paragraph(f"Missing image: {resolved_image_path}")
            continue

        if line.startswith("#"):
            level = min(len(line) - len(line.lstrip("#")), 4)
            heading_text = line[level:].strip()
            if heading_text:
                docx_doc.add_heading(heading_text, level=level)
            continue

        if line.startswith("*LLM image note:*"):
            paragraph = docx_doc.add_paragraph()
            run = paragraph.add_run("LLM image note: ")
            run.bold = True
            paragraph.add_run(line.replace("*LLM image note:*", "", 1).strip())
            continue

        docx_doc.add_paragraph(line)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    docx_doc.save(output_docx)
    LOG.info("Saved DOCX generated from augmented markdown to %s", output_docx)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="One-shot Docling pipeline with LLM-ready Markdown output.")
    parser.add_argument(
        "input_file",
        type=Path,
        help="Path to the file to process (.pdf, .docx, .pptx, .xlsx).",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=Path("outputs/pipeline"),
        help="Folder for JSON, markdown, and extracted images.",
    )
    parser.add_argument(
        "--artifacts-path",
        type=Path,
        default=Path("models/docling"),
        help="Local Docling model artifacts directory.",
    )
    parser.add_argument(
        "--ocr-engine",
        choices=["rapidocr", "ocrmac", "easyocr", "none", "auto"],
        default="rapidocr",
        help="OCR engine for scanned PDFs. Use none to disable.",
    )
    parser.add_argument(
        "--device",
        choices=["auto", "cpu", "mps", "cuda", "xpu"],
        default="auto",
        help="Inference device.",
    )
    parser.add_argument(
        "--image-scale",
        type=float,
        default=2.0,
        help="Scaling factor for exported images. Higher values improve clarity but increase processing time and memory.",
    )
    parser.add_argument("--no-tables", action="store_true", help="Disable table structure extraction.")
    parser.add_argument("--verbose", action="store_true", help="Enable debug logging.")
    parser.add_argument(
        "--docx-output",
        type=Path,
        help="Optional path for generated DOCX with embedded images and stub LLM notes. Defaults to <output-dir>/read_llm.docx",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    configure_logging(args.verbose)
    start_time = time.perf_counter()
    try:
        if not args.input_file.exists():
            raise FileNotFoundError(f"Input file not found: {args.input_file}")
        LOG.info(
            "Processing %s with ocr=%s, image_scale=%.2f",
            args.input_file,
            args.ocr_engine,
            args.image_scale,
        )
        doc, md_path, _image_dir = run_docling(args)
        llm_stub = build_llm_stub(md_path)
        augmented_md_path, inserted_notes = augment_markdown_with_llm(md_path)

        docx_path = args.docx_output or args.output_dir / f"{args.input_file.stem}_read_llm.docx"
        generate_docx_from_augmented_markdown(
            augmented_md_path,
            docx_path,
            title=doc.name,
        )

        LOG.info(
            "Pipeline complete. Markdown ready at %s, augmented markdown at %s, and DOCX at %s",
            md_path,
            augmented_md_path,
            docx_path,
        )
        LOG.info("Total processing time: %.2f seconds", time.perf_counter() - start_time)
        LOG.debug("Augmented %s image tags", len(inserted_notes))
        LOG.debug("LLM stub response: %s", llm_stub)
    except Exception as exc:  # noqa: BLE001
        LOG.info("Processing stopped after %.2f seconds", time.perf_counter() - start_time)
        LOG.exception("Pipeline failed: %s", exc)
        raise


if __name__ == "__main__":
    main()
