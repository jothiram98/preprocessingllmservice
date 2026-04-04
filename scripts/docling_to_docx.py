from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docling_core.types.doc.document import DocItem, DoclingDocument, PictureItem, RefItem, TextItem
from docx import Document
from docx.shared import Inches


def build_picture_description(pic: PictureItem, doc: DoclingDocument) -> str:
    text_segments: list[str] = []
    for child in pic.children:
        item = child.resolve(doc)
        if isinstance(item, TextItem):
            clean = item.text.strip()
            if clean:
                text_segments.append(clean)

    caption_text = pic.caption_text(doc)
    if caption_text:
        text_segments.insert(0, caption_text.strip())

    return " ".join(text_segments).strip()


def add_item_to_docx(docx_doc: Document, item: DocItem, doc: DoclingDocument, max_width: float) -> None:
    if isinstance(item, TextItem):
        text = item.text.strip()
        if text:
            docx_doc.add_paragraph(text)
    elif isinstance(item, PictureItem):
        description = build_picture_description(item, doc)
        if description:
            para = docx_doc.add_paragraph()
            run = para.add_run("Screenshot description: ")
            run.bold = True
            para.add_run(description)

        image = item.get_image(doc)
        if image is not None:
            buffer = BytesIO()
            image.save(buffer, format="PNG")
            buffer.seek(0)
            docx_doc.add_picture(buffer, width=Inches(max_width))
            buffer.close()


def assemble_docx(doc_json: Path, output_docx: Path, max_image_width: float) -> None:
    doc = DoclingDocument.model_validate_json(doc_json.read_text(encoding="utf-8"))
    docx_doc = Document()
    if doc.name:
        docx_doc.add_heading(doc.name, level=1)

    for ref in doc.body.children:
        item = ref.resolve(doc)
        add_item_to_docx(docx_doc, item, doc, max_image_width)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    docx_doc.save(output_docx)
    print(f"Written DOCX to {output_docx}")


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Emit a DOCX with embedded image descriptions.")
    parser.add_argument(
        "document_json",
        type=Path,
        help="Docling JSON export (the `_embedded.json` output).",
    )
    parser.add_argument(
        "--output",
        "-o",
        type=Path,
        default=Path("outputs/docling_document.docx"),
        help="Target DOCX file.",
    )
    parser.add_argument(
        "--image-width",
        type=float,
        default=5.0,
        help="Max image width in inches.",
    )

    args = parser.parse_args()
    assemble_docx(args.document_json, args.output, args.image_width)
