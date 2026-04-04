from __future__ import annotations

import argparse
import base64
import os
from io import BytesIO
from pathlib import Path
from typing import Iterable

from docling_core.types.doc.document import (
    DoclingDocument,
    GroupItem,
    PictureItem,
    RefItem,
    TextItem,
)

from docx import Document
from dotenv import load_dotenv
from PIL import Image
import base64
import requests

load_dotenv()
_HF_TOKEN = os.getenv("HUGGINGFACE_API_TOKEN")
if not _HF_TOKEN:
    raise RuntimeError("HUGGINGFACE_API_TOKEN missing; please add it to .env.")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Build a DOCX with placeholders for images plus Groq descriptions."
    )
    parser.add_argument(
        "document_json",
        type=Path,
        help="Docling JSON export path (the `_embedded.json` output).",
    )
    parser.add_argument(
        "--output",
        "-o",
        type=Path,
        default=Path("outputs/docling_with_images.docx"),
        help="Target DOCX file to write.",
    )
    parser.add_argument(
        "--image-dir",
        type=Path,
        default=Path("outputs/extracted_images"),
        help="Where to save extracted image files.",
    )
    return parser.parse_args()


def describe_image_with_hf(image_path: Path) -> str:
    with Image.open(image_path) as image:
        image.thumbnail((400, 400))
        buffer = BytesIO()
        image.save(buffer, format="JPEG", quality=75)
        buffer.seek(0)
        encoded = base64.b64encode(buffer.read()).decode("ascii")

    payload = {
        "inputs": {
            "image": f"data:image/jpeg;base64,{encoded}",
            "instruction": "Describe the contents of this screenshot in a concise sentence.",
        }
    }
    headers = {
        "Authorization": f"Bearer {_HF_TOKEN}",
        "Accept": "application/json",
    }
    resp = requests.post(
        "https://api-inference.huggingface.co/models/Qwen/Qwen2.5-VL-7B-Instruct",
        headers=headers,
        json=payload,
        timeout=60,
    )
    resp.raise_for_status()
    data = resp.json()
    if isinstance(data, dict) and "error" in data:
        raise RuntimeError(f"Hugging Face error: {data['error']}")

    if isinstance(data, list) and data and isinstance(data[0], dict):
        return data[0].get("generated_text", "").strip()
    return str(data).strip()


def walk_items(refs: Iterable[RefItem], doc: DoclingDocument):
    for ref in refs:
        item = ref.resolve(doc)
        if isinstance(item, GroupItem):
            yield from walk_items(item.children, doc)
        else:
            yield item


def save_picture(pic: PictureItem, doc: DoclingDocument, image_dir: Path, index: int) -> Path | None:
    image = pic.get_image(doc)
    if image is None:
        return None
    image_dir.mkdir(parents=True, exist_ok=True)
    target = image_dir / f"picture-{index}.png"
    image.save(target, format="PNG")
    return target


def add_item_to_document(docx_doc: Document, item: TextItem) -> None:
    if isinstance(item, TextItem):
        text = item.text.strip()
        if text:
            docx_doc.add_paragraph(text)


def main() -> None:
    args = parse_args()
    doc = DoclingDocument.model_validate_json(
        args.document_json.read_text(encoding="utf-8")
    )
    docx_doc = Document()
    if doc.name:
        docx_doc.add_heading(doc.name, level=1)

    image_index = 0
    for item in walk_items(doc.body.children, doc):
        if isinstance(item, PictureItem):
            image_index += 1
            image_path = save_picture(item, doc, args.image_dir, image_index)
            description = ""
            if image_path is not None:
                description = describe_image_with_hf(image_path)
            placeholder = docx_doc.add_paragraph()
            placeholder.add_run(f"Image: {image_path or 'missing'}").bold = True
            if description:
                docx_doc.add_paragraph(f"Description: {description}")
        else:
            add_item_to_document(docx_doc, item)

    args.output.parent.mkdir(parents=True, exist_ok=True)
    docx_doc.save(args.output)
    print(f"Written DOCX to {args.output} with {image_index} image placeholders.")


if __name__ == "__main__":
    main()
